from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import shutil
import regex as re

class PageLayoutProcessor:
    def __init__(self, input_path, output_path):
        self.input_path = input_path
        self.output_path = output_path
        self.temp_dir = tempfile.mkdtemp()
        self.uses_nirmala_ui = False
        
    def process(self):
        """Process the document's page layout"""
        try:
            # Create a backup of the input file
            temp_input = os.path.join(self.temp_dir, 'temp_input.docx')
            shutil.copy2(self.input_path, temp_input)
            
            # Load the document
            doc = Document(temp_input)
            
            # Check if document uses Nirmala UI
            self._check_for_nirmala_ui(doc)
            
            # Process fonts throughout the document
            self._process_fonts(doc)
            
            # Remove contact details and URLs
            self._remove_contact_details(doc)
            
            # Remove price-related lines
            self._remove_price_related_lines(doc)
            
            # Process each section in the document
            for section in doc.sections:
                # 1. Set orientation to portrait
                self._set_portrait_orientation(section)
                
                # 2. Set all margins to 1 inch
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                
                # 3. Remove headers
                if section.header:
                    section.header.is_linked_to_previous = True
                    for paragraph in section.header.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)
                
                # 4. Remove footers
                if section.footer:
                    section.footer.is_linked_to_previous = True
                    for paragraph in section.footer.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)
                
                # 5. Remove page borders
                try:
                    section._sectPr.remove_all("w:pgBorders")
                except:
                    pass

                # 6. Check and convert multiple columns to single column
                self._convert_to_single_column(section)
                
                # 7. Remove watermark
                self._remove_watermark(section)
            
            # 8. Set line spacing for all text elements
            self._set_line_spacing(doc)
            
            # 9. Process images - convert to Picture (U) and center align
            self._process_images(doc)
            
            # Save to temporary file first
            temp_output = os.path.join(self.temp_dir, 'temp_output.docx')
            doc.save(temp_output)
            
            # Then copy to final destination
            shutil.copy2(temp_output, self.output_path)
            print(f"Successfully processed document layout: {self.output_path}")
            
        except Exception as e:
            print(f"Error processing document: {str(e)}")
        finally:
            self._cleanup_temp_files()

    def _check_for_nirmala_ui(self, doc):
        """Check if document uses Nirmala UI font"""
        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.name == 'Nirmala UI':
                        self.uses_nirmala_ui = True
                        print("Document uses Nirmala UI font, will use Kalpurush for Bengali and English")
                        return
                        
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.font.name == 'Nirmala UI':
                                    self.uses_nirmala_ui = True
                                    print("Document uses Nirmala UI font, will use Kalpurush for Bengali and English")
                                    return
        except Exception as e:
            print(f"Error checking for Nirmala UI: {str(e)}")

    def _process_fonts(self, doc):
        """Process fonts throughout the document"""
        try:
            # Process main document paragraphs
            for paragraph in doc.paragraphs:
                self._process_paragraph_fonts(paragraph)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._process_paragraph_fonts(paragraph)
            
            print("Successfully processed fonts")
        except Exception as e:
            print(f"Error processing fonts: {str(e)}")

    def _process_paragraph_fonts(self, paragraph):
        """Process fonts in a paragraph"""
        try:
            for run in paragraph.runs:
                if not run.text.strip():  # Skip empty runs
                    continue
                    
                script_type = self._detect_script(run.text)
                current_font = run.font.name if run.font.name else ''
                
                # Print debugging information
                print(f"Text: {run.text[:20]}...")
                print(f"Detected script: {script_type}")
                print(f"Current font: {current_font}")
                
                # Case 1: If document uses Nirmala UI
                if self.uses_nirmala_ui:
                    if script_type in ['bengali', 'english']:
                        if current_font != 'Kalpurush':
                            run.font.name = 'Kalpurush'
                            print(f"Converted '{run.text[:20]}...' from {current_font} to Kalpurush (Nirmala UI case)")
                    elif script_type == 'arabic':
                        if current_font.lower() not in ['al majeed quranic', 'almajeedquranic']:
                            run.font.name = 'Al Majeed Quranic'
                            print(f"Converted Arabic text '{run.text[:20]}...' from {current_font} to Al Majeed Quranic")
                
                # Case 2: Normal case - only change if font doesn't match requirements
                else:
                    # For Bengali text
                    if script_type == 'bengali':
                        if current_font.lower() not in ['sutonnymj', 'sutonny mj']:
                            run.font.name = 'SutonnyMJ'
                            print(f"Converted Bengali text '{run.text[:20]}...' from {current_font} to SutonnyMJ")
                        else:
                            print(f"Keeping existing SutonnyMJ font for: {run.text[:20]}...")
                    
                    # For English text
                    elif script_type == 'english':
                        if current_font.lower() not in ['times new roman', 'timesnewroman']:
                            run.font.name = 'Times New Roman'
                            print(f"Converted English text '{run.text[:20]}...' from {current_font} to Times New Roman")
                        else:
                            print(f"Keeping existing Times New Roman font for: {run.text[:20]}...")
                    
                    # For Arabic text
                    elif script_type == 'arabic':
                        if current_font.lower() not in ['al majeed quranic', 'almajeedquranic']:
                            run.font.name = 'Al Majeed Quranic'
                            print(f"Converted Arabic text '{run.text[:20]}...' from {current_font} to Al Majeed Quranic")
                        else:
                            print(f"Keeping existing Al Majeed Quranic font for: {run.text[:20]}...")
                
        except Exception as e:
            print(f"Error processing paragraph fonts: {str(e)}")

    def _detect_script(self, text):
        """Detect the script type of the text"""
        try:
            # Common SutonnyMJ Bengali characters and combinations
            sutonnymj_chars = {
                # Basic Bengali characters in SutonnyMJ
                'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z',
                'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                # Special characters and modifiers
                '†', 'Š', '‡', 'ÿ', '¨', '©', '®', '¯', '°', '±', '²', '³', '´', 'µ', '¶', '·', '¸', '¹', 'º', '»', '¼', '½', '¾', '¿',
                # Additional Bengali characters
                'È', 'É', 'Ê', 'Ë', 'Ì', 'Í', 'Î', 'Ï', 'Ð', '����', 'Ò', 'Ó', 'Ô', 'Õ', 'Ö', '×', 'Ø', 'Ù', 'Ú', 'Û', 'Ü', 'Ý', 'Þ', 'ß',
                # Common combinations
                'ww', 'vv', 'šš', '††', '‡‡', 'ii', 'yy'
            }
            
            # Count SutonnyMJ Bengali characters in the text
            bengali_char_count = sum(1 for char in text if char in sutonnymj_chars)
            total_length = len(text.strip())
            
            # If text is empty, return english
            if total_length == 0:
                return 'english'
            
            # Calculate percentage of Bengali characters
            bengali_percentage = (bengali_char_count / total_length) * 100
            
            # Check for common SutonnyMJ patterns
            has_bengali_markers = any(marker in text for marker in ['†', 'Š', '‡', 'ÿ', '©', '®'])
            
            # If text has Bengali markers or significant Bengali characters, consider it Bengali
            if has_bengali_markers or bengali_percentage > 10:
                return 'bengali'
            
            # Check for Unicode Bengali (fallback)
            bengali_unicode = re.compile(r'[\u0980-\u09FF\u200C\u200D]')
            if bengali_unicode.search(text):
                return 'bengali'
            
            # Check for Arabic
            arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]')
            if arabic_pattern.search(text):
                return 'arabic'
            
            # Check if the text looks like pure English
            english_pattern = re.compile(r'^[a-zA-Z0-9\s\.,!?\'\"()-]*$')
            if english_pattern.match(text):
                return 'english'
            
            # If we still have some Bengali characters, consider it Bengali
            if bengali_char_count > 0:
                return 'bengali'
            
            # Default to English if nothing else matches
            return 'english'
            
        except Exception as e:
            print(f"Error detecting script: {str(e)}")
            return 'english'  # Default to English on error

    def _process_images(self, doc):
        """Process all images in the document - convert to Picture (U) and center align"""
        try:
            image_count = 0
            
            # Process images in main document body
            for paragraph in doc.paragraphs:
                if self._has_image(paragraph):
                    self._process_paragraph_images(paragraph)
                    image_count += 1
                    print(f"Processed image {image_count}")
            
            # Process images in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if self._has_image(paragraph):
                                self._process_paragraph_images(paragraph)
                                image_count += 1
                                print(f"Processed image {image_count}")
            
            print(f"Successfully processed {image_count} images")
        except Exception as e:
            print(f"Error processing images: {str(e)}")

    def _has_image(self, paragraph):
        """Check if paragraph contains an image"""
        try:
            for run in paragraph.runs:
                if len(run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'})) > 0:
                    return True
            return False
        except Exception as e:
            print(f"Error checking for images: {str(e)}")
            return False

    def _process_paragraph_images(self, paragraph):
        """Process images in a paragraph while maintaining their position"""
        try:
            # Center align the paragraph containing the image
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process each run in the paragraph
            for run in paragraph.runs:
                # Find all pictures in the run
                pics = run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'})
                
                for pic in pics:
                    try:
                        # Get the parent drawing element
                        drawing = pic.getparent().getparent()
                        
                        # Check if it's inline or floating
                        if drawing.tag.endswith('}inline'):
                            # For inline images, just ensure paragraph is centered
                            continue
                            
                        elif drawing.tag.endswith('}anchor'):
                            # For floating images, set position to center
                            pos_h = drawing.find('.//wp:positionH', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                            if pos_h is not None:
                                pos_h.set('relativeFrom', 'margin')
                                align = pos_h.find('.//wp:align', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                                if align is not None:
                                    align.text = 'center'
                                
                    except Exception as e:
                        print(f"Error processing picture: {str(e)}")
                        continue
                        
        except Exception as e:
            print(f"Error processing paragraph images: {str(e)}")

    def _remove_watermark(self, section):
        """Remove watermark from the document"""
        try:
            # Remove watermark from section properties
            if hasattr(section._sectPr, 'get_or_add_background'):
                # Remove background (which might contain watermark)
                section._sectPr.remove_all('w:background')
                
            # Remove any VML drawings (often used for watermarks)
            try:
                for child in section._sectPr.findall('.//{urn:schemas-microsoft-com:vml}shape'):
                    child.getparent().remove(child)
            except:
                pass
                
            # Remove any picture watermarks
            try:
                for child in section._sectPr.findall('.//w:pict'):
                    child.getparent().remove(child)
            except:
                pass
                
            # Remove header references that might contain watermarks
            try:
                for child in section._sectPr.findall('.//w:headerReference'):
                    child.getparent().remove(child)
            except:
                pass
                
            # Remove document background
            try:
                section._sectPr.remove_all('w:documentBackground')
            except:
                pass
                
            print("Successfully removed watermark")
        except Exception as e:
            print(f"Error removing watermark: {str(e)}")

    def _convert_to_single_column(self, section):
        """Check and convert multiple columns to single column"""
        try:
            # Get the columns element
            cols = section._sectPr.xpath("./w:cols")[0]
            
            # Check if document has multiple columns
            num_cols = int(cols.get(qn('w:num'))) if cols.get(qn('w:num')) else 1
            
            if num_cols > 1:
                print(f"Converting from {num_cols} columns to single column")
                # Set to single column
                cols.set(qn('w:num'), '1')
                
                # Remove any column spacing
                if cols.get(qn('w:space')):
                    cols.set(qn('w:space'), '0')
                    
                # Remove any specific column width settings
                for child in cols.getchildren():
                    cols.remove(child)
        except Exception as e:
            print(f"Error checking/converting columns: {str(e)}")

    def _set_line_spacing(self, doc):
        """Set line spacing to 1.15 for all text elements in the document"""
        try:
            # Set spacing for all styles in the document
            for style in doc.styles:
                if hasattr(style, 'paragraph_format'):
                    style.paragraph_format.line_spacing = 1.15

            # Process all paragraphs in main document
            for paragraph in doc.paragraphs:
                self._apply_spacing_to_paragraph(paragraph)
            
            # Process paragraphs in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._apply_spacing_to_paragraph(paragraph)
            
            # Process headers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self._apply_spacing_to_paragraph(paragraph)
                        
                # Process footers
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self._apply_spacing_to_paragraph(paragraph)
            
            print("Successfully set line spacing to 1.15 for all text elements")
        except Exception as e:
            print(f"Error setting line spacing: {str(e)}")

    def _apply_spacing_to_paragraph(self, paragraph):
        """Apply 1.15 line spacing to a paragraph"""
        if paragraph._element is not None:
            # Set line spacing at paragraph format level
            paragraph.paragraph_format.line_spacing = 1.15
            
            # Ensure the spacing is applied at the XML level
            if paragraph._p.pPr is None:
                paragraph._p.get_or_add_pPr()
            
            # Set spacing in XML
            spacing = paragraph._p.pPr.xpath('./w:spacing')
            if not spacing:
                spacing_element = OxmlElement('w:spacing')
                spacing_element.set(qn('w:line'), str(int(240 * 1.15)))  # 240 twips = 1 line
                spacing_element.set(qn('w:lineRule'), 'auto')
                paragraph._p.pPr.append(spacing_element)
            else:
                spacing[0].set(qn('w:line'), str(int(240 * 1.15)))
                spacing[0].set(qn('w:lineRule'), 'auto')

    def _cleanup_temp_files(self):
        """Clean up temporary files and directories"""
        try:
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                print("Cleaned up temporary files")
        except Exception as e:
            print(f"Error cleaning up temporary files: {str(e)}")

    def _set_portrait_orientation(self, section):
        """Set section orientation to portrait"""
        try:
            # Get section properties
            section_properties = section._sectPr
            
            # Create or get page size element
            page_size = section_properties.xpath('./w:pgSz')
            if not page_size:
                page_size = OxmlElement('w:pgSz')
                section_properties.append(page_size)
            else:
                page_size = page_size[0]
            
            # Set orientation to portrait
            page_size.set(qn('w:orient'), 'portrait')
            
            # Set standard page width and height for portrait (in twips)
            # A4 in portrait: 8.3 x 11.7 inches (or 11906 x 16838 twips)
            page_size.set(qn('w:w'), '11906')
            page_size.set(qn('w:h'), '16838')
            
            print("Set orientation to portrait")
        except Exception as e:
            print(f"Error setting portrait orientation: {str(e)}")

    def _remove_contact_details(self, doc):
        """Remove contact details and URLs from the document"""
        try:
            # Patterns to match
            patterns = {
                'phone': [
                    # Bangladesh mobile numbers (11 digits)
                    r'(?<!\d)01\d{3}[-\s]?\d{6}(?!\d)',  # 01XXX-XXXXXX
                    r'(?<!\d)01\d{2}[-\s]?\d{7}(?!\d)',  # 01XX-XXXXXXX
                    r'(?<!\d)01\d{1}[-\s]?\d{8}(?!\d)',  # 01X-XXXXXXXX
                    r'(?<!\d)01\d{9}(?!\d)',             # 01XXXXXXXXX
                    
                    # Hyphenated formats
                    r'(?<!\d)01\d{3}[-]\d{3}[-]\d{3}(?!\d)',  # 01XXX-XXX-XXX
                    r'(?<!\d)01\d{2}[-]\d{3}[-]\d{4}(?!\d)',  # 01XX-XXX-XXXX
                    r'(?<!\d)01\d{4}[-]\d{6}(?!\d)',          # 01XXXX-XXXXXX
                    
                    # With country code
                    r'(?<!\d)\+8801\d{9}(?!\d)',              # +8801XXXXXXXXX
                    r'(?<!\d)8801\d{9}(?!\d)',                # 8801XXXXXXXXX
                    
                    # Bengali digits (০১, ০২, etc.)
                    r'(?<!\d)[০১]\d{4}[-\s]?[০-৯\d]{6}(?!\d)',
                    r'(?<!\d)[০১][০-৯\d]{9}(?!\d)',
                    r'(?<!\d)[+৮৮][০১][০-৯\d]{9}(?!\d)',
                    
                    # With labels (both English and Bengali)
                    r'(?:Phone|Mobile|Contact|Tel|Call|ফোন|মোবাইল|কল|যোগাযোগ)[\s\:]+[\+\d\-\s০-৯]{8,}',
                    
                    # Specific format like your example
                    r'(?<!\d)0\d{4}[-]0?\d{5,6}(?!\d)',      # 01568-069216
                    
                    # General mobile number patterns
                    r'(?<!\d)(?:013|014|015|016|017|018|019)\d{8}(?!\d)',
                    r'(?<!\d)(?:০১৩|০১৪|০১৫|০১৬|০১৭|০১৮|০১৯)[০-৯]{8}(?!\d)'
                ],
                'url': r'(?:https?://|www\.)\S+\.(?:com|org|net|bd|edu|gov|info|biz)',
                'email': r'[\w\-\.]+@[\w\-\.]+\.[a-zA-Z]{2,}',
                'social': r'(?:facebook\.com|fb\.com|twitter\.com|linkedin\.com|instagram\.com)/[\w\-\.]+',
                'whatsapp': r'(?:whatsapp|viber|imo)[\s\:]+\d[\d\-\s]{9,}'
            }
            
            # Function to check if text is a date
            def is_date(text):
                date_patterns = [
                    r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}',  # DD/MM/YYYY or similar
                    r'\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',
                    r'\d{4}[-/]\d{1,2}[-/]\d{1,2}',  # YYYY/MM/DD or similar
                    r'(?:19|20)\d{2}[-/]\d{1,2}[-/]\d{1,2}'  # Year first date format
                ]
                return any(re.search(pattern, text, re.IGNORECASE) for pattern in date_patterns)
            
            # Function to check if text is ISBN
            def is_isbn(text):
                isbn_patterns = [
                    r'ISBN\s*[:।-]?\s*(?:\d+[- ]?){4,}',  # ISBN followed by digits
                    r'ISBN\s*\d*\s*[:।-]?\s*(?:\d+[- ]?){4,}',  # ISBN with optional number
                    r'(?:\d{3}-\d{3}-\d{4}-\d{2}-\d{1})',  # 13-digit ISBN format
                    r'(?:\d{3}-\d{1}-\d{4}-\d{4}-\d{1})'   # Alternative ISBN format
                ]
                return any(re.search(pattern, text, re.IGNORECASE) for pattern in isbn_patterns)
            
            # Function to check if text is a reference number
            def is_reference_number(text):
                ref_patterns = [
                    r'No\.\s*\d+/[A-Z]+/\d+/\d+',  # Common reference number format
                    r'[A-Z]+/\d+/[A-Z]+/\d+'  # Alternative reference format
                ]
                return any(re.search(pattern, text) for pattern in ref_patterns)
            
            # Function to check if text contains contact info
            def contains_contact_info(text):
                if is_isbn(text) or is_date(text) or is_reference_number(text):
                    return False
                    
                for pattern_list in patterns.values():
                    if isinstance(pattern_list, list):
                        for pattern in pattern_list:
                            if re.search(pattern, text, re.IGNORECASE):
                                return True
                    elif re.search(pattern_list, text, re.IGNORECASE):
                        return True
                return False
            
            # Function to remove only contact info from text while preserving other content
            def remove_contact_info(paragraph):
                if any(is_isbn(run.text) or is_date(run.text) or is_reference_number(run.text) for run in paragraph.runs):
                    return False
                    
                has_changes = False
                for run in paragraph.runs:
                    if is_isbn(run.text) or is_date(run.text) or is_reference_number(run.text):
                        continue
                        
                    new_text = run.text
                    for pattern_list in patterns.values():
                        if isinstance(pattern_list, list):
                            for pattern in pattern_list:
                                matches = list(re.finditer(pattern, new_text, re.IGNORECASE))
                                if matches:
                                    has_changes = True
                                    for match in reversed(matches):
                                        start, end = match.span()
                                        new_text = new_text[:start] + new_text[end:]
                        else:
                            matches = list(re.finditer(pattern_list, new_text, re.IGNORECASE))
                            if matches:
                                has_changes = True
                                for match in reversed(matches):
                                    start, end = match.span()
                                    new_text = new_text[:start] + new_text[end:]
                    
                    if has_changes:
                        run.text = new_text.strip()
                
                return has_changes
            
            # Process main document paragraphs
            for paragraph in doc.paragraphs:
                if contains_contact_info(paragraph.text):
                    print(f"Processing paragraph with contact info: {paragraph.text.strip()}")
                    remove_contact_info(paragraph)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if contains_contact_info(paragraph.text):
                                print(f"Processing table cell with contact info: {paragraph.text.strip()}")
                                remove_contact_info(paragraph)
            
            print("Successfully removed contact details and URLs")
        except Exception as e:
            print(f"Error removing contact details: {str(e)}")

    def _remove_price_related_lines(self, doc):
        """Remove price-related lines from the document"""
        try:
            # Patterns to match price-related text in Bengali (SutonnyMJ encoding) and English
            price_patterns = [
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)মূল্য\s*[:।].*',  # মূল্য followed by : or । and any text
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)g~j¨\s*[:।].*',  # মূল্য in SutonnyMJ
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)দাম\s*[:।].*',   # দাম followed by : or । and any text
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)`vg\s*[:।].*',   # দাম in SutonnyMJ
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)টাকা\s*মাত্র',    # টাকা মাত্র
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)UvKv\s*gvÎ',     # টাকা মাত্র in SutonnyMJ
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)মূল্যঃ',         # মূল্যঃ
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)g~j¨t',         # মূল্যঃ in SutonnyMJ
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)Price\s*:(?!.*ISBN).*',   # Price: in English
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)\d+\s*টাকা(?!\s*ISBN)',     # Number followed by টাকা
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)\d+\s*UvKv(?!\s*ISBN)',    # Number followed by টাকা in SutonnyMJ
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)Taka\s*:(?!.*ISBN).*',    # Taka: in English
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)(?:TK|Tk)\s*\.?\s*\d+(?!\s*ISBN)', # TK/Tk followed by number
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)৳\s*\d+(?!\s*ISBN)',       # Bengali Taka symbol followed by number
                r'(?<!ISBN\s*[:।-]\s*)(?<!ISBN\s*\d*\s*[:।-]\s*)\.{3,}\s*\d+\s*(?:টাকা|UvKv|Taka|TK|Tk)?(?!\s*ISBN)', # Dots followed by price
            ]
            
            # Function to check if text is ISBN
            def is_isbn(text):
                isbn_patterns = [
                    r'ISBN\s*[:।-]?\s*\d[\d\-]*',
                    r'ISBN\s*[:।-]?\s*\d[\d\s\-]*\d',
                    r'ISBN\s*\d*\s*[:।-]?\s*\d[\d\-]*'
                ]
                return any(re.search(pattern, text, re.IGNORECASE) for pattern in isbn_patterns)
            
            # Compile patterns
            patterns = [re.compile(pattern, re.IGNORECASE) for pattern in price_patterns]
            
            # Function to check if text contains price-related information
            def contains_price_info(text):
                if is_isbn(text):  # Skip if text contains ISBN
                    return False
                return any(pattern.search(text) for pattern in patterns)
            
            # Process main document paragraphs
            paragraphs_to_remove = []
            for paragraph in doc.paragraphs:
                if contains_price_info(paragraph.text):
                    paragraphs_to_remove.append(paragraph._element)
                    print(f"Found price-related text to remove: {paragraph.text.strip()}")
            
            # Remove marked paragraphs
            for paragraph_element in paragraphs_to_remove:
                parent = paragraph_element.getparent()
                if parent is not None:
                    parent.remove(paragraph_element)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs_to_remove = []
                        for paragraph in cell.paragraphs:
                            if contains_price_info(paragraph.text):
                                paragraphs_to_remove.append(paragraph._element)
                                print(f"Found price-related text to remove in table: {paragraph.text.strip()}")
                        
                        # Remove marked paragraphs from cell
                        for paragraph_element in paragraphs_to_remove:
                            parent = paragraph_element.getparent()
                            if parent is not None:
                                parent.remove(paragraph_element)
            
            print("Successfully removed price-related lines")
        except Exception as e:
            print(f"Error removing price-related lines: {str(e)}")

def main():
    # Update these paths according to your file locations
    input_file = "/Users/macbookpro/Desktop/assignment_rokomari/Project eBook Automation/Ebook/90191.docx"
    output_file = "output.docx"
    
    # Process the document
    processor = PageLayoutProcessor(input_file, output_file)
    processor.process()

if __name__ == "__main__":
    main() 